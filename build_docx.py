#!/usr/bin/env python3
"""Render resume.md into a styled .docx.

Purpose-built for this resume's markdown subset: headings, bullets, tables,
horizontal rules, and inline **bold** / *italic* / [links](url).

    uv run --with python-docx python build_docx.py [resume.md] [Scott_Long_Resume.docx]
"""

from __future__ import annotations

import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ACCENT = RGBColor(0x1A, 0x3D, 0x5C)  # deep slate blue for headings
MUTED = RGBColor(0x44, 0x4A, 0x52)  # body-adjacent grey for meta lines
BODY = RGBColor(0x1A, 0x1A, 0x1A)
FONT = "Calibri"

# **bold** | *italic* | [text](url) | plain
INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|\[[^\]]+?\]\([^)]+?\))")


def add_hyperlink(paragraph, url: str, text: str, size: float) -> None:
    """python-docx has no link API; build the w:hyperlink element directly."""
    r_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    for tag, val in (("w:color", "1A5FB4"), ("w:u", "single")):
        el = OxmlElement(tag)
        el.set(qn("w:val"), "single" if tag == "w:u" else val)
        props.append(el)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    props.append(fonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    props.append(sz)
    run.append(props)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    link.append(run)
    paragraph._p.append(link)


def write_inline(paragraph, text: str, size: float = 9.5, color: RGBColor = BODY, bold: bool = False) -> None:
    """Emit `text` into `paragraph`, honouring bold/italic/link markup."""
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            run, run.bold = paragraph.add_run(tok[2:-2]), True
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            run, run.italic = paragraph.add_run(tok[1:-1]), True
        elif tok.startswith("["):
            link = re.match(r"\[([^\]]+)\]\(([^)]+)\)", tok)
            if link:
                add_hyperlink(paragraph, link.group(2), link.group(1), size)
                continue
            run = paragraph.add_run(tok)
        else:
            run = paragraph.add_run(tok)
        run.font.name = FONT
        run.font.size = Pt(size)
        run.font.color.rgb = color
        if bold:
            run.bold = True


def tight(paragraph, before: float = 0, after: float = 0, spacing: float = 1.06) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = spacing


def rule(doc) -> None:
    """A thin horizontal divider, drawn as a bottom border on an empty paragraph."""
    p = doc.add_paragraph()
    tight(p, before=2, after=6)
    pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B8C2CC")
    borders.append(bottom)
    pr.append(borders)


def build(src: str, dest: str) -> None:
    lines = open(src, encoding="utf-8").read().splitlines()

    doc = Document()
    doc.styles["Normal"].font.name = FONT
    doc.styles["Normal"].font.size = Pt(9.5)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(0.5)
        section.left_margin = section.right_margin = Inches(0.62)

    i, first_h2 = 0, True
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        # --- horizontal rule ------------------------------------------------
        if stripped == "---":
            i += 1
            continue

        # --- table ----------------------------------------------------------
        if stripped.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{2,}:?", c or "-") for c in cells):
                    rows.append(cells)
                i += 1
            rows = [r for r in rows if any(c for c in r)]
            if rows:
                width = max(len(r) for r in rows)
                table = doc.add_table(rows=0, cols=width)
                table.alignment = WD_TABLE_ALIGNMENT.LEFT
                table.autofit = False
                for row in rows:
                    cells = table.add_row().cells
                    for n, text in enumerate(row):
                        cell = cells[n]
                        cell.width = Inches(1.28 if n == 0 else 6.0)
                        p = cell.paragraphs[0]
                        tight(p, before=1.5, after=1.5)
                        write_inline(p, text, size=9)
            continue

        # --- headings -------------------------------------------------------
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()

            if level == 1:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                tight(p, after=1)
                run = p.add_run(text)
                run.font.name = FONT
                run.font.size = Pt(23)
                run.bold = True
                run.font.color.rgb = ACCENT
                pr = run._element.get_or_add_rPr()
                spacing = OxmlElement("w:spacing")
                spacing.set(qn("w:val"), "24")  # slight letterspacing on the name
                pr.append(spacing)

            elif level == 2:
                p = doc.add_paragraph()
                tight(p, before=0 if first_h2 else 9, after=3)
                run = p.add_run(text.upper())
                run.font.name = FONT
                run.font.size = Pt(10.5)
                run.bold = True
                run.font.color.rgb = ACCENT
                pr = run._element.get_or_add_rPr()
                spacing = OxmlElement("w:spacing")
                spacing.set(qn("w:val"), "16")
                pr.append(spacing)
                first_h2 = False
                # underline the section label
                bpr = p._p.get_or_add_pPr()
                borders = OxmlElement("w:pBdr")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "8")
                bottom.set(qn("w:space"), "2")
                bottom.set(qn("w:color"), "1A3D5C")
                borders.append(bottom)
                bpr.append(borders)

            elif level == 3:
                p = doc.add_paragraph()
                tight(p, before=8, after=1)
                write_inline(p, text, size=11.5, color=ACCENT, bold=True)

            else:  # level 4 — initiative sub-headers
                p = doc.add_paragraph()
                tight(p, before=7, after=2)
                write_inline(p, text, size=10, color=RGBColor(0x2C, 0x2C, 0x2C), bold=True)
            i += 1
            continue

        # --- bullets --------------------------------------------------------
        if stripped.startswith("- "):
            body = stripped[2:]
            i += 1
            while i < len(lines) and lines[i].startswith("  ") and lines[i].strip() and not lines[i].strip().startswith("- "):
                body += " " + lines[i].strip()
                i += 1
            p = doc.add_paragraph(style="List Bullet")
            tight(p, before=1.5, after=1.5, spacing=1.08)
            p.paragraph_format.left_indent = Inches(0.21)
            p.paragraph_format.first_line_indent = Inches(-0.13)
            write_inline(p, body)
            continue

        # --- blank ----------------------------------------------------------
        if not stripped:
            i += 1
            continue

        # --- paragraph (soft-wrapped: gather the block) ----------------------
        block = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(r"^(#|-\s|\||---)", lines[i].strip()):
            block.append(lines[i].strip())
            i += 1

        # A run of short italic/meta lines under a heading stays line-per-line.
        meta = all(ln.startswith("*") and ln.endswith("*") for ln in block)
        contact = any("@" in ln or "LinkedIn" in ln for ln in block)

        if meta or contact:
            for ln in block:
                p = doc.add_paragraph()
                tight(p, after=1)
                if contact:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                write_inline(p, ln, size=9 if contact else 9.5, color=MUTED)
        else:
            p = doc.add_paragraph()
            tight(p, after=4, spacing=1.1)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            write_inline(p, " ".join(block))

    doc.save(dest)
    print(f"wrote {dest}")


if __name__ == "__main__":
    src = sys.argv[1] if len(sys.argv) > 1 else "resume.md"
    dest = sys.argv[2] if len(sys.argv) > 2 else "Scott_Long_Resume.docx"
    build(src, dest)
